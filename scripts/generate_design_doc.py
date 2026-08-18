"""
scripts/generate_design_doc.py  —  v4  (complete redesign)
------------------------------------------------------------
Generates notes/MCP_Hub_Design_Document.docx

Design principles:
  - 6 focused sections (not 14): Overview, Auth, Authz, Hub, MCP, Ops
  - Minimal template: 2 accent colours, 3 callout types, no phase banners
  - One diagram anchors each section; code illustrates, does not define
  - Consistent section format: lead paragraph → visual → key details
"""

import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import diagrams as D

OUT  = Path(__file__).parent.parent / "notes" / "MCP_Hub_Design_Document.docx"

# ---------------------------------------------------------------------------
# Palette — deliberately minimal: 2 primaries + 2 accents + neutrals
# ---------------------------------------------------------------------------
C = {
    "navy":   "1F3864",   # H1, table headers, emphasis
    "blue":   "2563EB",   # H2 accent, code border, bullet markers
    "dk":     "1A3F80",   # H1 underline rule, cover rule
    "grey":   "374151",   # body text
    "mgrey":  "6B7280",   # muted / captions / footer
    "lgrey":  "F3F4F6",   # code bg, table alt rows, note bg
    "white":  "FFFFFF",
    "amber":  "B45309",   # warning accent text
    "ltamb":  "FFFBEB",   # warning background
    "purple": "5B21B6",   # security accent text
    "ltpur":  "F5F3FF",   # security background
    "ltblue": "EFF6FF",   # note background
    "bdr":    "D1D5DB",   # light border
}


def rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for o in tcPr.findall(qn("w:shd")):
        tcPr.remove(o)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _cell_borders(cell, color="D1D5DB", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        t = OxmlElement(f"w:{side}")
        t.set(qn("w:val"), "single")
        t.set(qn("w:sz"), sz)
        t.set(qn("w:space"), "0")
        t.set(qn("w:color"), color)
        tcB.append(t)
    tcPr.append(tcB)


def _para_border_bottom(p, color="2563EB", sz="12"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _run(para, text, bold=False, italic=False, size=11,
         color="374151", font="Calibri", underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    return r


def _set_page_margins(doc, t=2.0, b=2.0, l=2.0, r=2.0):
    sec = doc.sections[0]
    sec.top_margin    = Cm(t)
    sec.bottom_margin = Cm(b)
    sec.left_margin   = Cm(l)
    sec.right_margin  = Cm(r)


# ---------------------------------------------------------------------------
# Typography — clean 3-level hierarchy
# ---------------------------------------------------------------------------

def H1(doc, text):
    """Major section — navy 18 pt bold + thick blue underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(8)
    _para_border_bottom(p, C["dk"], "16")
    _run(p, text, bold=True, size=18, color=C["navy"])
    return p


def H2(doc, text):
    """Sub-section — blue 13 pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, text, bold=True, size=13, color=C["blue"])
    return p


def H3(doc, text):
    """Tertiary — navy 11.5 pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    _run(p, text, bold=True, size=11.5, color=C["navy"])
    return p


def lead(doc, text):
    """Intro paragraph — sets context for a section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    _run(p, text, size=11, color=C["grey"])
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, size=11, color=C["grey"])
    return p


def bullet(doc, text, label=None, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.left_indent       = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    mark = "▸" if level else "•"
    _run(p, f"{mark}  ", bold=True, size=10.5, color=C["blue"])
    if label:
        _run(p, label, bold=True, size=11, color=C["navy"])
    _run(p, text, size=11, color=C["grey"])
    return p


def code(doc, text, caption=None):
    """Code block — Consolas 8.5 pt on light grey with blue left accent."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _cell_bg(cell, C["lgrey"])
    # blue left bar, light borders on other sides
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side, clr, sz in [
        ("left",   C["blue"],  "20"),
        ("top",    C["bdr"],   "4"),
        ("bottom", C["bdr"],   "4"),
        ("right",  C["bdr"],   "4"),
    ]:
        t = OxmlElement(f"w:{side}")
        t.set(qn("w:val"), "single")
        t.set(qn("w:sz"), sz)
        t.set(qn("w:space"), "0")
        t.set(qn("w:color"), clr)
        tcB.append(t)
    tcPr.append(tcB)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.35)
    r = p.add_run(text)
    r.font.name  = "Consolas"
    r.font.size  = Pt(8.5)
    r.font.color.rgb = rgb("1E293B")
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(8)
        cp.paragraph_format.left_indent  = Cm(0.35)
        _run(cp, caption, italic=True, size=9, color=C["mgrey"])
    else:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)
    return tbl


def callout(doc, kind, title, text):
    """Callout — 3 kinds only: note (blue), warning (amber), security (purple)."""
    cfg = {
        "note":     (C["ltblue"], C["blue"],   "NOTE"),
        "warning":  (C["ltamb"],  C["amber"],  "WARNING"),
        "security": (C["ltpur"],  C["purple"], "SECURITY"),
    }
    bg, accent, lbl = cfg.get(kind, cfg["note"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _cell_bg(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side, clr, sz in [
        ("left",   accent,    "20"),
        ("top",    C["bdr"],  "4"),
        ("bottom", C["bdr"],  "4"),
        ("right",  C["bdr"],  "4"),
    ]:
        t = OxmlElement(f"w:{side}")
        t.set(qn("w:val"), "single")
        t.set(qn("w:sz"), sz)
        t.set(qn("w:space"), "0")
        t.set(qn("w:color"), clr)
        tcB.append(t)
    tcPr.append(tcB)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.35)
    r1 = p.add_run(lbl)
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = rgb(accent)
    if title:
        r2 = p.add_run(f"  —  {title}\n")
        r2.bold = True
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = rgb(accent)
    r3 = p.add_run(text)
    r3.font.name = "Calibri"
    r3.font.size = Pt(10)
    r3.font.color.rgb = rgb(C["grey"])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl


def table(doc, headers, rows, widths=None, stripe=True):
    """Data table — navy header, clean thin borders, alternating row shading."""
    cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _cell_bg(cell, C["navy"])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.15)
        _run(p, h, bold=True, size=9.5, color=C["white"])

    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = C["lgrey"] if (stripe and ri % 2 == 0) else C["white"]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            _cell_bg(cell, bg)
            _cell_borders(cell, C["bdr"], "4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.15)
            _run(p, str(val), size=9.5, color=C["grey"])

    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


def img(doc, path, width_cm=16.0, caption=None):
    if not path or not Path(str(path)).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(12)
        _run(cp, caption, italic=True, size=9, color=C["mgrey"])
    return p


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    _para_border_bottom(p, C["bdr"], "4")
    return p


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _cover(doc):
    spacer(doc, 60)

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t1, "MCP Hub", bold=True, size=38, color=C["navy"])

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(4)
    _run(t2, "Solution Design Document", bold=False, size=20, color=C["blue"])

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(16)
    _para_border_bottom(rule, C["dk"], "24")

    # Metadata table — centred
    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Version",    "4.0"),
        ("Status",     "CONFIDENTIAL — Client Distribution"),
        ("Date",       datetime.date.today().strftime("%B %d, %Y")),
        ("Covers",     "Architecture · Authentication · Authorization · Operations"),
    ]
    for ri, (lbl, val) in enumerate(data):
        lc = meta_tbl.rows[ri].cells[0]
        vc = meta_tbl.rows[ri].cells[1]
        _cell_bg(lc, C["navy"])
        _cell_bg(vc, C["lgrey"])
        pl = lc.paragraphs[0]
        pv = vc.paragraphs[0]
        for p in (pl, pv):
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.left_indent  = Cm(0.2)
        _run(pl, lbl, bold=True, size=10, color=C["white"])
        _run(pv, val, size=10.5, color=C["grey"])
    meta_tbl.columns[0].width = Cm(3.0)
    meta_tbl.columns[1].width = Cm(13.0)


# ---------------------------------------------------------------------------
# Executive Summary
# ---------------------------------------------------------------------------

def _exec_summary(doc):
    doc.add_page_break()
    H1(doc, "Executive Summary")

    body(doc,
         "The MCP Hub is a zero-trust control plane for AI agent-to-service authentication. "
         "It provides a single gateway that centralises server discovery, cryptographic "
         "authentication, intelligent routing, and structured audit — eliminating bespoke "
         "per-agent integration code across all backend services.")

    body(doc,
         "An agent presents a signed hub JWT to POST /discover. The hub validates it, runs "
         "an LLM-guided routing step to select the best-fit MCP server, and mints a "
         "short-lived scoped token for it. The agent calls the MCP server directly over "
         "Streamable HTTP using that token — the hub is not in the data path for tool calls.")

    spacer(doc, 4)
    table(doc,
          ["Design Decision", "Rationale"],
          [
              ("RS256 asymmetric JWT signing",
               "Private key stays in the hub. MCP servers receive only the public key and can "
               "verify tokens — but cannot forge them. A compromised MCP server cannot impersonate the hub."),
              ("Per-server audience scoping (aud = server_id)",
               "Each JWT is bound to one server. A stolen token is unusable on any other server. "
               "Blast radius = 1 server × 1 hour."),
              ("Multi-provider auth (local / Azure AD)",
               "Local mode: RS256 hub tokens, HS256 chat tokens, or static API key. "
               "Azure mode: Entra ID OIDC — no local keys required. Switch via AUTH_PROVIDER env var."),
              ("LangGraph ReAct routing agent",
               "New instance per /discover request prevents concurrent state collision. "
               "LLM selects best server by capability; falls back to first active server if unavailable."),
              ("4-way observability fan-out",
               "Every event written to in-memory deque, stdout, JSONL file, and MySQL simultaneously. "
               "MySQL failure permanently disables that sink — other three remain active."),
          ],
          widths=[5.5, 10.5])


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def _toc(doc):
    doc.add_page_break()
    H1(doc, "Table of Contents")
    spacer(doc, 4)

    entries = [
        ("1",   "System Architecture",                      True),
        ("1.1", "Architecture Overview",                    False),
        ("1.2", "Component Roles",                          False),
        ("1.3", "Layer Communication Reference",            False),
        ("2",   "Authentication",                           True),
        ("2.1", "Auth Provider Configuration",              False),
        ("2.2", "Token Verification Flow",                  False),
        ("2.3", "JWT Token Lifecycle & Key Infrastructure", False),
        ("2.4", "Server Discovery & Token Scoping",         False),
        ("2.5", "MCP Server JWT Validation",                False),
        ("3",   "Authorization",                            True),
        ("3.1", "Role Model",                               False),
        ("3.2", "Per-Tool RBAC Enforcement",                False),
        ("3.3", "Credential Isolation",                     False),
        ("4",   "Hub Service",                              True),
        ("4.1", "Server Registry",                          False),
        ("4.2", "LLM-Based Routing",                        False),
        ("4.3", "Observability & Audit",                    False),
        ("4.4", "Discovery API & Event Model",              False),
        ("5",   "MCP Integration",                          True),
        ("5.1", "Streamable HTTP Transport",                False),
        ("5.2", "Chat Service Integration",                 False),
        ("6",   "Operations & Deployment",                  True),
        ("6.1", "Environment Variables",                    False),
        ("6.2", "Production Readiness Checklist",           False),
        ("6.3", "Security Reference",                       False),
        ("6.4", "Key Files Reference",                      False),
    ]

    tbl = doc.add_table(rows=len(entries), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (num, title, is_h1) in enumerate(entries):
        lc = tbl.rows[ri].cells[0]
        rc = tbl.rows[ri].cells[1]
        bg = C["lgrey"] if is_h1 else C["white"]
        _cell_bg(lc, bg)
        _cell_bg(rc, bg)
        for cell in (lc, rc):
            _cell_borders(cell, C["bdr"], "2")
        p1 = lc.paragraphs[0]
        p2 = rc.paragraphs[0]
        for p in (p1, p2):
            p.paragraph_format.space_before = Pt(4 if is_h1 else 2)
            p.paragraph_format.space_after  = Pt(4 if is_h1 else 2)
            p.paragraph_format.left_indent  = Cm(0.1 if is_h1 else 0.5)
        sz  = 10 if is_h1 else 9.5
        clr = C["navy"] if is_h1 else C["grey"]
        _run(p1, num,   bold=is_h1, size=sz, color=clr)
        _run(p2, title, bold=is_h1, size=sz, color=clr)
    tbl.columns[0].width = Cm(1.6)
    tbl.columns[1].width = Cm(14.4)


# ---------------------------------------------------------------------------
# 1. System Architecture
# ---------------------------------------------------------------------------

def _s1(doc, paths):
    doc.add_page_break()
    H1(doc, "1.  System Architecture")
    lead(doc,
         "The MCP Hub sits between AI agents and backend services. It handles authentication, "
         "server discovery, and routing so that agents can use any registered MCP server "
         "without knowing its endpoint, credentials, or capabilities in advance.")

    H2(doc, "1.1  Architecture Overview")
    img(doc, paths.get("system"), width_cm=16.0,
        caption="Figure 1 — System Overview: Layers, Components, and Data Flows")

    H2(doc, "1.2  Component Roles")
    table(doc,
          ["Component", "Role", "Port"],
          [
              ("Chat UI (SPA)",      "Browser query interface — submits queries, displays streaming responses",       "—"),
              ("Admin UI (SPA)",     "Hub management — server registry, event logs, health, CRUD",                  "—"),
              ("Chat Server",        "User auth (PBKDF2-SHA256); session management; response streaming to browser", "8080"),
              ("MCP Hub Server",     "Server registry, JWT issuance, LLM routing, JWKS, event logging",            "8090"),
              ("Agent Orchestrator", "LangGraph ReAct loop; calls /discover, opens MCP sessions, invokes tools",    "—"),
              ("MCP Servers",        "FastMCP servers exposing domain tools over Streamable HTTP with JWT auth",     "9100+"),
              ("MySQL",              "Registry (mcp_servers), events (hub_events), conversations, users",           "3306"),
          ],
          widths=[4.0, 10.0, 2.0])

    callout(doc, "note", "Transport Protocol",
            "All agent-to-MCP-server communication uses Streamable HTTP (POST /mcp). "
            "Every request carries an Authorization: Bearer <token> header — not just the initial handshake. "
            "The session is tracked via an Mcp-Session-Id header returned during initialize.")

    H2(doc, "1.3  Layer Communication Reference")
    body(doc,
         "Each boundary in the stack has a fixed HTTP contract. The table below lists the "
         "method, mandatory headers, request body shape, and response shape for every layer "
         "transition. No credential or header passes through more than one boundary.")

    table(doc,
          ["Boundary", "Method & URL", "Request", "Response"],
          [
              ("Browser → Chat Server",
               "POST /login",
               'Body: {"username":"alice","password":"..."}\nContent-Type: application/json',
               '{"access_token":"<hub JWT>","token_type":"bearer","sub":"alice","roles":["agent"]}\nSets HttpOnly cookie: session=<token>'),
              ("Browser → Chat Server",
               "POST /messages",
               'Cookie: session=<hub JWT>\nBody: {"query":"..."}\nContent-Type: application/json',
               '200 OK (task created); events streamed via GET /sse or polled via GET /poll'),
              ("Agent → Hub Server",
               "POST /auth/login",
               'Body: {"username":"...","password":"..."}\nContent-Type: application/json',
               '{"access_token":"<RS256 JWT>","token_type":"bearer","sub":"...","roles":[...]}'),
              ("Agent → Hub Server",
               "POST /discover",
               'Authorization: Bearer <hub JWT>\nContent-Type: application/json\nBody: {"intent":"customer C007 loyalty status"}',
               '{"servers":[{config + "server_token":"<RS256 JWT aud=server_id>"}],"method":"agent","reason":"...","hub_metadata":{...},"auth_meta":{...}}'),
              ("Agent → MCP Server",
               "POST /mcp  (initialize)",
               'Authorization: Bearer <server JWT>\nAccept: application/json, text/event-stream\nContent-Type: application/json\nBody: JSON-RPC initialize',
               'Header: Mcp-Session-Id: <uuid>\nBody: {"result":{"protocolVersion":"2024-11-05","capabilities":{},"serverInfo":{"name":"..."}}}'),
              ("Agent → MCP Server",
               "POST /mcp  (tools/call)",
               'Authorization: Bearer <server JWT>\nMcp-Session-Id: <uuid>\nBody: JSON-RPC tools/call with name + arguments',
               '{"result":{"content":[{"type":"text","text":"{...result JSON...}"}]}}'),
              ("MCP Server → MySQL",
               "SQL query (SQLAlchemy)",
               'Credentials: MYSQL_USER / MYSQL_PASSWORD\nParameterised query against semantic view',
               'Result rows from customer_360_view, pricing_view, etc.'),
          ],
          widths=[2.8, 3.4, 5.6, 4.2])

    callout(doc, "note", "Credential isolation at every boundary",
            "Authorization: Bearer in the agent→MCP request carries the per-server JWT — "
            "not the hub JWT or any MySQL credential. "
            "MYSQL_USER/MYSQL_PASSWORD are never forwarded beyond the MCP server process boundary.")


# ---------------------------------------------------------------------------
# 2. Authentication
# ---------------------------------------------------------------------------

def _s2(doc, paths):
    doc.add_page_break()
    H1(doc, "2.  Authentication")
    lead(doc,
         "Authentication is centralised in hub_service/auth.py. "
         "It supports multiple providers (local RS256/HS256, Azure AD/Entra ID, static API key) "
         "and enforces a strict verification priority chain on every hub API call.")

    img(doc, paths.get("auth"), width_cm=16.0,
        caption="Figure 2 — End-to-End Authentication Flow")

    H2(doc, "2.1  Auth Provider Configuration")
    body(doc,
         "Two top-level flags in hub_service/auth.py govern all auth behaviour. "
         "AUTH_PROVIDER selects the trust anchor. AUTH_ENABLED can disable all checks "
         "for local development only — never set false in production.")

    table(doc,
          ["AUTH_PROVIDER", "Trust Anchor", "Tokens Accepted", "When to Use"],
          [
              ("local (default)", "Hub RSA private key",
               "RS256 JWT (hub-minted) · HS256 JWT (JWT_SECRET) · static HUB_API_KEY",
               "Self-hosted; no external identity provider"),
              ("azure",           "Azure AD / Entra ID JWKS",
               "Microsoft-issued OIDC tokens (RS256)",
               "Enterprise; existing Azure AD tenant"),
              ("(disabled)",      "None — AUTH_ENABLED=false",
               "All requests pass as anonymous admin",
               "Local dev only — never staging/production"),
          ],
          widths=[2.8, 3.5, 6.2, 3.5])

    code(doc,
"# hub_service/auth.py — loaded at module import\n"
"AUTH_ENABLED  = os.environ.get('AUTH_ENABLED', 'true').lower() in ('1','true','yes')\n"
"AUTH_PROVIDER = os.environ.get('AUTH_PROVIDER', 'local')  # 'local' | 'azure'\n\n"
"# local provider\n"
"HUB_API_KEY  = os.environ.get('HUB_API_KEY', '')    # static pre-shared key\n"
"JWT_SECRET   = os.environ.get('JWT_SECRET', '')     # HS256 shared secret\n"
"# azure provider\n"
"AZURE_TENANT_ID = os.environ.get('AZURE_TENANT_ID', '')\n"
"AZURE_CLIENT_ID = os.environ.get('AZURE_CLIENT_ID', '')")

    H2(doc, "2.2  Token Verification Flow")
    body(doc,
         "verify_token() implements a strict priority chain. Steps are tried in order "
         "and the first match wins. Algorithm routing (Step 3) reads the JWT header "
         "without signature verification to detect RS256 vs HS256 before choosing the path.")

    code(doc,
"# verify_token() priority chain — hub_service/auth.py\n\n"
"Step 0  AUTH_ENABLED = false      → pass, claims = {sub: anonymous, roles: [admin]}\n"
"Step 1  AUTH_PROVIDER = azure     → _verify_azure()        # OIDC via Microsoft JWKS\n"
"Step 2  token == HUB_API_KEY      → pass, claims = {sub: api-key-user, roles: [...]}\n"
"Step 3  JWT header alg = RS256    → _verify_jwt_rs256()    # sig · iss · aud · exp\n"
"Step 4  JWT_SECRET is configured  → _verify_jwt_local()    # HS256 sig · exp\n"
"Step 5  RS256 last resort         → _verify_jwt_rs256()    # unknown alg, try anyway\n"
"Step 6  _DEV_MODE_ACTIVE          → pass, claims = {sub: dev, roles: [admin]}\n"
"Step 7  No match                  → 401 Unauthorized")

    callout(doc, "warning", "Dev Mode — Exact Activation Conditions",
            "_DEV_MODE_ACTIVE fires only when ALL four hold simultaneously:\n"
            "  1.  AUTH_ENABLED = true\n"
            "  2.  AUTH_PROVIDER = local\n"
            "  3.  HUB_API_KEY is not set\n"
            "  4.  JWT_SECRET is not set\n\n"
            "Hub emits a ⚠ startup WARNING. Setting JWT_SECRET is the minimum to exit dev mode.\n"
            "RSA key files are auto-generated on startup but do NOT by themselves exit dev mode.")

    H2(doc, "2.3  JWT Token Lifecycle & Key Infrastructure")
    img(doc, paths.get("jwt"), width_cm=16.0,
        caption="Figure 3 — JWT Token Lifecycle & RSA Key Infrastructure")

    body(doc,
         "The hub generates a 2048-bit RSA key pair on first startup "
         "(hub_service/.keys/private.pem and public.pem). "
         "The private key signs all tokens and never leaves the hub process. "
         "The public key is published at GET /.well-known/jwks.json with a kid (Key ID) "
         "header that enables key rotation without downtime.")

    code(doc,
"# JWKS endpoint — GET /.well-known/jwks.json\n"
'{ "keys": [{ "kty":"RSA", "kid":"hub-rsa-1", "use":"sig", "alg":"RS256",\n'
'             "n":"<base64url modulus>", "e":"AQAB" }] }\n\n'
"# Token minting CLI — development and testing\n"
"python hub_service/auth.py --sub agent --roles agent --hours 24\n"
"python hub_service/auth.py --sub admin --roles admin,agent --hours 8\n\n"
"# Key rotation: update HUB_JWT_KID → restart hub → wait 1h for old tokens to expire")

    H2(doc, "2.4  Server Discovery & Token Scoping")
    body(doc,
         "When an agent calls POST /discover, the hub validates the caller's hub JWT, "
         "runs LLM-based routing to select the best-fit server, then mints a separate "
         "short-lived JWT for each matched server. Each token is bound exclusively to its target:")

    code(doc,
"# hub_service/auth.py generate_server_token()\n"
"payload = {\n"
"    'iss':       'fab-mcp-hub',    # must match MCP_JWT_ISSUER on MCP server\n"
"    'aud':       server_id,        # only this server accepts this token\n"
"    'sub':       caller_sub,       # forwarded from hub JWT — never elevated\n"
"    'roles':     caller_roles,     # forwarded from hub JWT — never elevated\n"
"    'server_id': server_id,        # additional claim for MCP server filtering\n"
"    'exp':       now + 3600,       # 1 hour — shorter than hub session (8h)\n"
"}")

    callout(doc, "security", "Why Per-Server Audience Scoping",
            "Each JWT carries aud = server_id. Server B rejects a token intended for Server A "
            "because the audience claim does not match.\n"
            "Blast radius of any single token compromise: 1 server × 1 hour.\n"
            "A single token valid on all servers would create unlimited cross-service blast radius.")

    H2(doc, "2.5  MCP Server JWT Validation")
    body(doc,
         "MCP servers run two middleware layers in sequence on every Streamable HTTP request. "
         "FastMCP's JWTVerifier handles cryptographic validation; "
         "BearerClaimsMiddleware decodes the already-verified claims into a per-request ContextVar "
         "for RBAC — without a second JWKS round-trip.")

    code(doc,
"# mcp_server/server.py — middleware registration (order matters)\n\n"
"# Layer 1 — FastMCP JWTVerifier (cryptographic validation)\n"
"verifier = JWTVerifier(\n"
"    jwks_uri = f'{HUB_URL}/.well-known/jwks.json',   # fetches RSA public key from hub\n"
"    issuer   = os.environ['MCP_JWT_ISSUER'],          # must match HUB_JWT_ISSUER\n"
"    audience = os.environ['MCP_SERVER_ID']            # must match aud in token\n"
")  # returns HTTP 401 on: invalid sig | wrong aud | wrong iss | expired\n\n"
"# Layer 2 — BearerClaimsMiddleware (decode for RBAC — no 2nd JWKS call)\n"
"class BearerClaimsMiddleware(BaseHTTPMiddleware):\n"
"    async def dispatch(self, request, call_next):\n"
"        token = request.headers.get('Authorization','').removeprefix('Bearer ').strip()\n"
"        payload = jwt.decode(token, options={'verify_signature': False})\n"
"        _request_claims.set({'sub': payload['sub'], 'roles': payload.get('roles',[])})\n"
"        return await call_next(request)")


# ---------------------------------------------------------------------------
# 3. Authorization
# ---------------------------------------------------------------------------

def _s3(doc, paths):
    doc.add_page_break()
    H1(doc, "3.  Authorization")
    lead(doc,
         "Authorization operates at two levels: hub-side RBAC (role checked on every hub API call) "
         "and per-tool RBAC inside each MCP server (role enforced before any business logic executes).")

    H2(doc, "3.1  Role Model")
    table(doc,
          ["Role", "Hub Access", "MCP Tool Access", "Typical Holder"],
          [
              ("admin",
               "All endpoints — /discover · /servers · /health · /api/hub/* (full CRUD)",
               "All tools — bypasses role checks (admin is implicit in require_role)",
               "Operator / admin login"),
              ("agent",
               "POST /discover · GET /servers · GET /health",
               "Tools annotated require_role('agent') or require_role('agent','admin')",
               "AI agent / orchestrator process"),
              ("readonly",
               "GET /servers · GET /health only",
               "Read-only tools only",
               "Monitoring system / audit consumer"),
          ],
          widths=[2.0, 5.0, 5.0, 4.0])

    callout(doc, "note", "Role Forwarding — No Elevation",
            "Roles from the hub JWT are copied unchanged into every per-server JWT — never elevated. "
            "An agent presenting an 'agent' hub token always gets 'agent' server tokens. "
            "The hub cannot be tricked into issuing elevated server tokens.")

    H2(doc, "3.2  Per-Tool RBAC Enforcement")
    body(doc,
         "Every tool function calls require_role() as its first statement — before any I/O, "
         "database access, or input validation. Claims come from BearerClaimsMiddleware via ContextVar.")

    code(doc,
"# Standard tool pattern — mcp_server/tools.py\n"
"@mcp.tool()\n"
"def customer_360(customer_id: str) -> dict:\n"
"    require_role('agent', 'admin')               # 1st: before any I/O\n"
"    audit_log('customer_360',                    # 2nd: logs keys NOT values (PII)\n"
"              args={'customer_id': customer_id})\n"
"    if not customer_id or len(customer_id) > 50:\n"
"        raise ValueError('customer_id: 1-50 chars')\n"
"    return db.query_one(\n"
"        'SELECT * FROM customer_360_view WHERE id = %s',\n"
"        (customer_id,)    # parameterised query — never string-format user input\n"
"    )")

    img(doc, paths.get("rbac"), width_cm=14.5,
        caption="Figure 4 — Per-Tool RBAC Decision Flow")

    H2(doc, "3.3  Credential Isolation")
    img(doc, paths.get("credentials"), width_cm=16.0,
        caption="Figure 5 — Credential Isolation: No Credential Crosses a Layer Boundary")

    table(doc,
          ["Layer", "Credential Used", "Valid Scope", "Never Forwarded To"],
          [
              ("Browser → Chat Server",
               "Hub JWT (8h, HttpOnly cookie)",
               "Chat server session only",
               "MCP servers · MySQL · external APIs"),
              ("Agent → Hub",
               "Hub JWT (8h, Authorization header)",
               "/discover · /servers · /health",
               "MCP servers · MySQL · external APIs"),
              ("Agent → MCP Server",
               "Per-server JWT (1h, aud = server_id)",
               "Single MCP server only",
               "Hub · MySQL · other MCP servers"),
              ("MCP Server → MySQL",
               "MYSQL_USER / MYSQL_PASSWORD",
               "DB queries inside MCP server",
               "Agent · Hub · browser"),
              ("MCP Server → External APIs",
               "MCP_TOOL_KEY (env var)",
               "Tool-specific HTTP calls",
               "Agent · Hub · browser"),
          ],
          widths=[3.5, 4.0, 3.5, 5.0])


# ---------------------------------------------------------------------------
# 4. Hub Service
# ---------------------------------------------------------------------------

def _s4(doc, paths):
    doc.add_page_break()
    H1(doc, "4.  Hub Service")
    lead(doc,
         "The hub (hub_service/hub_server.py) is the operational core. "
         "It owns the server registry, issues JWTs, routes discovery requests via an LLM agent, "
         "and records every event to four simultaneous sinks.")

    H2(doc, "4.1  Server Registry")
    body(doc,
         "The registry is a MySQL table (mcp_servers) seeded from mcp-hub.json via seed_hub_db.py. "
         "Servers can be added, edited, disabled, or probed through the admin UI or REST API. "
         "A 60-second in-process cache minimises MySQL round-trips; "
         "POST /api/hub/refresh busts it immediately.")

    code(doc,
"# hub_service/mcp-hub.json — declarative source (seeded into MySQL)\n"
'{\n'
'  "servers": [{\n'
'    "id":          "fab-customer-server",\n'
'    "endpoint":    "http://mcp-customer:9100/mcp",\n'
'    "transport":   "streamable-http",\n'
'    "capability":  "customer data — 360 view, order history, loyalty",\n'
'    "skills":      ["customer lookup", "order history", "loyalty points"],\n'
'    "is_active":   true\n'
'  }]\n'
'}')

    table(doc,
          ["Endpoint", "Auth", "Purpose"],
          [
              ("GET  /health",               "Public",        "Service status, hub metadata, registered server count"),
              ("GET  /.well-known/jwks.json","Public",        "RSA public key in JWK Set format for MCP server verification"),
              ("POST /auth/login",           "Public",        "username+password → RS256 JWT (8h); entry point for agent auth"),
              ("POST /discover",             "agent | admin", "LLM routing → server selection → per-server JWT minting"),
              ("GET  /servers",              "agent | admin", "Full active server list with capabilities and skills"),
              ("GET  /api/logs",             "admin",         "Structured event log — MySQL primary, in-memory fallback"),
              ("POST /api/hub/refresh",      "admin",         "Bust 60-second registry cache immediately"),
              ("CRUD /api/hub/servers/*",    "admin",         "Add, edit, enable/disable, delete, probe server entries"),
          ],
          widths=[4.2, 2.4, 9.4])

    H2(doc, "4.2  LLM-Based Routing")
    body(doc,
         "Every POST /discover creates a fresh LangGraph ReAct agent instance. "
         "The agent receives all active servers' capability, skills, description, and examples, "
         "then reasons about which server best fits the query. "
         "A new instance per request is required — sharing one across concurrent calls corrupts internal message state.")

    code(doc,
"# hub_server.py — routing\n"
"def _agent_route(query, servers):\n"
"    tool  = _make_routing_tool(servers)               # stateless pick_server tool\n"
"    agent = create_react_agent(llm, [tool])           # fresh instance per request\n"
"    result = agent.invoke({\n"
"        'messages': [('user', routing_prompt(query, servers))]\n"
"    })                                                # LLM: THINK → CALL pick_server → OBSERVE\n"
"    return extract_selected_ids(result)")

    callout(doc, "note", "Routing Fallback",
            "If the LLM is unavailable (timeout, model not loaded, network error), "
            "the hub falls back to returning the first active server in the registry. "
            "Monitor routing.method='fallback' events in GET /api/logs to detect LLM availability issues.")

    H2(doc, "4.3  Observability & Audit")
    body(doc,
         "Every log_event() call fans out to four sinks simultaneously. "
         "MySQL is permanently disabled after the first failure — "
         "retrying on every event call would add latency to every HTTP request.")

    table(doc,
          ["Sink", "Access", "Failure Behaviour"],
          [
              ("In-memory deque (500 entries)", "GET /api/logs",         "Never fails — oldest entries evicted"),
              ("stdout",                        "docker logs / console", "Never fails"),
              ("logs/hub.log (JSONL)",          "tail / log aggregator", "Silently skips if directory not writable"),
              ("MySQL hub_events",              "GET /api/logs primary", "Permanently disabled on first error; restart hub to re-enable"),
          ],
          widths=[4.8, 3.5, 7.7])

    table(doc,
          ["Event Type", "When Emitted", "Key data Fields"],
          [
              ("auth",           "Every token check",       "valid · sub · roles · token_type · iss · provider · endpoint · method"),
              ("request",        "Every HTTP request",      "method · path · status · latency_ms"),
              ("routing",        "Server selection",         "method(agent|first_match) · server_ids · reason · intent · sub"),
              ("request_detail", "POST /discover handler",  "Full req/resp headers + body · auth_sub · auth_roles (verbose trace)"),
              ("admin",          "Server CRUD actions",     "action(create|update|delete_server) · server_id · changed_by"),
              ("error",          "Runtime exceptions",      "message · traceback"),
          ],
          widths=[2.5, 3.5, 10.0])

    H2(doc, "4.4  Discovery API & Event Model")
    body(doc,
         "Section 4.4 provides the complete data model: the mcp_servers registry schema, "
         "the hub_events audit table, the POST /discover request and response contracts, "
         "and an explanation of how each layer connects through the hub.")

    img(doc, paths.get("schema"), width_cm=16.0,
        caption="Figure 7 — Hub Database Schema: mcp_servers and hub_events tables")

    H3(doc, "mcp_servers — Registry Table Schema")
    body(doc,
         "The mcp_servers table is the single source of truth for all registered backends. "
         "It is seeded from hub_service/mcp-hub.json via scripts/seed_hub_db.py (UPSERT — safe to re-run). "
         "All fields except id, name, and endpoint have safe defaults; "
         "is_active=0 disables a server without deleting it.")

    table(doc,
          ["Column", "Type", "Purpose"],
          [
              ("id",              "VARCHAR(100)  PK",   "Unique server identifier. Used as JWT aud claim and MCP_SERVER_ID."),
              ("name",            "VARCHAR(255)",        "Human-readable display name shown in Admin UI and logs."),
              ("endpoint",        "VARCHAR(500)",        "Full HTTP URL for POST /mcp calls, e.g. http://mcp-customer:9100/mcp."),
              ("transport",       "VARCHAR(50)",         "'streamable-http' (FAB data servers) or 'sse' (demo servers). Controls which MCP client is used."),
              ("capability",      "TEXT",                "Natural-language summary of server capabilities. Included verbatim in the LLM routing prompt."),
              ("skills",          "JSON  (array)",       "Array of skill tags (e.g. ['customer lookup', 'order history']). Used by routing agent for semantic matching."),
              ("description",     "TEXT",                "Extended description providing additional routing context for edge cases."),
              ("examples",        "JSON  (array)",       "Array of sample query strings this server handles well. Helps LLM routing identify intent."),
              ("start_cmd",       "TEXT",                "Local development startup command. Informational only — not used at runtime."),
              ("api_key",         "VARCHAR(1000)",       "Optional static Bearer token. When set, overrides hub-minted JWT for agent→MCP authentication."),
              ("api_key_expires", "TIMESTAMP",           "Expiry for api_key. NULL means the key never expires. Hub does not enforce this — checked manually."),
              ("is_active",       "TINYINT(1)",          "1 = included in /discover responses. 0 = disabled and excluded, but row is retained for history."),
              ("created_at",      "TIMESTAMP",           "Auto-set on INSERT. Tracks when the server was registered."),
              ("updated_at",      "TIMESTAMP  ON UPDATE","Auto-updated on every change. Tracks last admin edit."),
          ],
          widths=[2.8, 2.6, 10.6])

    callout(doc, "note", "Changelog table",
            "Every CREATE / UPDATE / DELETE action on mcp_servers is also recorded in "
            "mcp_server_changelog (server_id · action · changed_by · before_state JSON · after_state JSON). "
            "This provides a full audit trail of registry changes without overwriting the live row.")

    H3(doc, "POST /discover — Request & Response")
    body(doc,
         "POST /discover is the primary integration point. "
         "It accepts a natural-language intent, routes it to the best-fit server, "
         "and returns a scoped JWT alongside all connection details the agent needs.")

    code(doc,
"# POST /discover — request body (Pydantic DiscoverRequest)\n"
'{ "intent": "what is the loyalty score for customer C007?" }\n\n'
"# Required headers\n"
"Authorization: Bearer <hub JWT>    # RS256, exp=8h, iss=fab-mcp-hub\n"
"Content-Type: application/json\n\n"
"# Response body (DiscoverResponse) — key fields\n"
"{\n"
'  "servers": [\n'
"    {\n"
'      "id":           "fab-customer-server",\n'
'      "name":         "FAB Customer Intelligence MCP Server",\n'
'      "endpoint":     "http://mcp-customer:9100/mcp",\n'
'      "transport":    "streamable-http",\n'
'      "capability":   "customer 360 data, order history, loyalty score",\n'
'      "skills":       ["customer lookup", "loyalty points", "order history"],\n'
'      "server_token": "<RS256 JWT — aud=fab-customer-server, exp=now+3600>"\n'
"    }\n"
"  ],\n"
'  "method":   "agent",          // "agent" = LLM routed | "first_match" = LLM unavailable\n'
'  "reason":   "Server handles customer 360 and loyalty queries",\n'
'  "hub_metadata": { "hub_name":"FAB MCP Hub", "version":"3.0", "server_count":2 },\n'
'  "auth_meta":    { "sub":"alice", "roles":["agent"], "token_type":"jwt", "exp":1720000000 }\n'
"}")

    H3(doc, "hub_events — Audit Table Schema")
    body(doc,
         "Every hub operation — authentication, routing, HTTP request, admin action, error — "
         "is written to hub_events as a structured JSON row. "
         "The data column stores the full event dict so any event type can be queried by field name "
         "using MySQL JSON path expressions (e.g. data->>'$.sub').")

    table(doc,
          ["Column", "Type", "Purpose"],
          [
              ("id",         "BIGINT  AUTO_INCREMENT  PK", "Monotonically increasing event ID."),
              ("ts",         "DOUBLE  (indexed)",           "Unix timestamp as float. Indexed (idx_ts) for range queries."),
              ("type",       "VARCHAR(64)  (indexed)",      "Event category — auth · request · routing · request_detail · admin · error. Indexed (idx_type)."),
              ("data",       "JSON",                        "Full event payload. Fields vary by type (see table below)."),
              ("created_at", "TIMESTAMP",                   "Wall-clock insert time. ts is always populated; created_at is the DB default."),
          ],
          widths=[2.8, 3.4, 9.8])

    table(doc,
          ["Event Type", "data Fields", "Notes"],
          [
              ("auth",
               "valid · sub · roles · token_type · iss · provider · endpoint · method · bearer_token",
               "Logged on every hub API call. token_type: jwt | apikey | dev."),
              ("request",
               "method · path · status · latency_ms",
               "Logged by _RequestLogMiddleware for every HTTP request."),
              ("routing",
               "method · server_ids · reason · intent (capped 120 chars) · sub",
               "method=agent (LLM) or first_match (fallback). server_ids is a list."),
              ("request_detail",
               "endpoint · request_headers · request_body · response_status · response_headers · response_body · auth_sub · auth_roles",
               "Verbose trace of /discover. Useful for debugging routing decisions."),
              ("admin",
               "action · server_id · changed_by",
               "Written on every server create/update/delete. action: create_server | update_server | delete_server."),
              ("error",
               "message · traceback",
               "Unhandled exceptions. Traceback is truncated if too long."),
          ],
          widths=[2.8, 7.2, 6.0])

    callout(doc, "warning", "tool_audit events — hub vs. MCP side",
            "tool_audit events (tool name, args_keys, sub, roles) are written by audit_log() "
            "in datalayer-as-service/mcp_server/auth.py — inside the MCP server process, "
            "not the hub. They are NOT written to hub_events. "
            "They appear in the MCP server's stdout only. "
            "To centralise MCP-side audit, configure log shipping from the MCP server container.")


# ---------------------------------------------------------------------------
# 5. MCP Integration
# ---------------------------------------------------------------------------

def _s5(doc, paths):
    doc.add_page_break()
    H1(doc, "5.  MCP Integration")
    lead(doc,
         "MCP servers expose domain tools over Streamable HTTP. "
         "The agent opens a session, attaches a scoped JWT on every request, "
         "and executes tool calls inside a LangGraph ReAct loop.")

    img(doc, paths.get("flow"), width_cm=16.0,
        caption="Figure 8 — End-to-End Request Flow: HTTP Contracts, Headers & Payloads at Each Layer Boundary")

    img(doc, paths.get("lifecycle"), width_cm=15.0,
        caption="Figure 9 — MCP Server Lifecycle: States and Transitions")

    H2(doc, "5.1  Streamable HTTP Transport")
    body(doc,
         "Streamable HTTP (POST /mcp) is the only MCP transport used in this stack. "
         "The agent negotiates a session ID during initialize and attaches it — along with "
         "the Authorization header — to every subsequent call.")

    code(doc,
"# agent.py mcp_session() — Streamable HTTP\n"
"async with streamablehttp_client(\n"
"    server['endpoint'],                    # e.g. http://mcp-customer:9100/mcp\n"
"    headers={'Authorization': f'Bearer {server_token}'}\n"
") as (read, write, _):\n"
"    async with ClientSession(read, write) as session:\n"
"        await session.initialize()         # POST /mcp → returns Mcp-Session-Id header\n"
"        tools  = await session.list_tools()\n"
"        result = await session.call_tool('customer_360', {'customer_id': 'C001'})")

    table(doc,
          ["Header", "Required", "Value / Notes"],
          [
              ("Content-Type",   "Yes", "application/json"),
              ("Accept",         "Yes", "application/json, text/event-stream — both MIME types required; FastMCP returns 406 if either is missing"),
              ("Authorization",  "Yes", "Bearer <per-server JWT> — present on EVERY request, not just initialize"),
              ("Mcp-Session-Id", "After initialize", "UUID returned in initialize response header; must be included on tools/list and tools/call"),
          ],
          widths=[3.2, 2.0, 10.8])

    H3(doc, "MCP JSON-RPC 2.0 Message Format")
    body(doc,
         "The Streamable HTTP transport carries MCP's JSON-RPC 2.0 protocol. "
         "There are three message types in the agent lifecycle: initialize (mandatory first call), "
         "tools/list (optional capability discovery), and tools/call (tool execution).")

    code(doc,
"# 1. initialize — mandatory first call on every new session\n"
'# POST /mcp   Authorization: Bearer <server JWT>   Accept: application/json, text/event-stream\n'
'{ "jsonrpc":"2.0", "id":"probe-init", "method":"initialize",\n'
'  "params": { "protocolVersion":"2024-11-05", "capabilities":{},\n'
'              "clientInfo": {"name":"mcp-client","version":"0.1"} } }\n\n'
"# Response — 200 OK + header  Mcp-Session-Id: <uuid>\n"
'{ "jsonrpc":"2.0", "result": {\n'
'    "protocolVersion":"2024-11-05", "capabilities": {"tools":{},"resources":{},"prompts":{}},\n'
'    "serverInfo": {"name":"FAB Customer Intelligence MCP Server","version":"1.0"}\n'
'  }, "id":"probe-init" }')

    code(doc,
"# 2. tools/list — discover available tools (optional; agent.py always calls this)\n"
'# POST /mcp   Mcp-Session-Id: <uuid>   Authorization: Bearer <server JWT>\n'
'{ "jsonrpc":"2.0", "id":2, "method":"tools/list", "params":{} }\n\n'
"# 3. tools/call — execute a tool\n"
'# POST /mcp   Mcp-Session-Id: <uuid>   Authorization: Bearer <server JWT>\n'
'{ "jsonrpc":"2.0", "id":"call-customer_360", "method":"tools/call",\n'
'  "params": { "name":"customer_360", "arguments": {"customer_id":"C007"} } }\n\n'
"# Response\n"
'{ "jsonrpc":"2.0",\n'
'  "result": { "content": [{"type":"text","text":"{\\"customer_id\\":\\"C007\\",...}"}] },\n'
'  "id":"call-customer_360" }')

    callout(doc, "note", "Session reuse within one agent loop",
            "A single agent invocation for one user query opens one MCP session (one initialize call). "
            "The LangGraph ReAct loop may call tools/call multiple times within that session — "
            "each carrying the same Mcp-Session-Id and a refreshed Authorization header. "
            "The session ends when the agent context manager exits.")

    callout(doc, "warning", "Token Renewal on 401",
            "When a 1-hour per-server JWT expires mid-session the MCP server returns HTTP 401.\n"
            "Renewal: (1) call POST /discover again with the hub JWT to mint a fresh server token, "
            "(2) open a new MCP session with the new token.\n"
            "Do not attempt to reuse the expired token or the existing Mcp-Session-Id — both are rejected.")

    H2(doc, "5.2  Chat Service Integration")
    body(doc,
         "The chat server (chat_service/chat_server.py) provides the browser interface. "
         "It authenticates users with PBKDF2-SHA256, holds hub JWTs in session cookies, "
         "and runs the agent as a background asyncio.Task so browser disconnects do not abort "
         "in-progress agent work. The final answer is always saved to MySQL.")

    table(doc,
          ["Endpoint", "Purpose", "Notes"],
          [
              ("POST /login",        "PBKDF2 verify → hub JWT (8h) as HttpOnly cookie",         "Rate-limited: 10 attempts/15 min per username → HTTP 423"),
              ("GET  /sse",          "Open long-lived streaming connection for response events",  "Background task survives browser disconnect"),
              ("POST /messages",     "Submit query → launch run_agent() asyncio.Task",           "Task created before SSE generator — decoupled lifetimes"),
              ("GET  /poll",         "Re-attach to running task; retrieve missed events",         "Used on reconnect after disconnect"),
              ("GET  /conversations","User conversation history from MySQL",                      "Paginated; requires valid session cookie"),
          ],
          widths=[3.5, 7.5, 5.0])


# ---------------------------------------------------------------------------
# 6. Operations & Deployment
# ---------------------------------------------------------------------------

def _s6(doc, paths):
    doc.add_page_break()
    H1(doc, "6.  Operations & Deployment")
    lead(doc,
         "Configuration reference, production readiness checklist, "
         "security properties summary, and key files index.")

    H2(doc, "6.1  Environment Variables")
    callout(doc, "note", "Startup-time variables",
            "Variables in the Hub Auth group are read by hub_service/auth.py at module import time. "
            "Changes require a hub process restart.")

    table(doc,
          ["Variable", "Service", "Required", "Description"],
          [
              # Hub auth
              ("AUTH_ENABLED",          "Hub Auth",   "No",       "true (default). false = all requests pass as anonymous admin — never production."),
              ("AUTH_PROVIDER",         "Hub Auth",   "No",       "local (default) or azure. local: RS256/HS256/API-key chain. azure: Entra ID OIDC."),
              ("HUB_API_KEY",           "Hub Auth",   "No",       "Static pre-shared key accepted as Bearer token. Roles from HUB_API_KEY_ROLES."),
              ("HUB_API_KEY_ROLES",     "Hub Auth",   "No",       "Comma-separated roles for HUB_API_KEY callers (default: agent)."),
              ("JWT_SECRET",            "Hub Auth",   "No",       "Shared HMAC secret for HS256 JWT verification. Setting this exits dev mode."),
              ("HUB_JWT_ISSUER",        "Hub Auth",   "No",       "iss claim in hub-minted tokens (default: fab-mcp-hub). Must match MCP_JWT_ISSUER."),
              ("HUB_JWT_KID",           "Hub Auth",   "No",       "kid in JWT header and JWKS (default: hub-rsa-1). Increment on RSA key rotation."),
              ("HUB_PRIVATE_KEY_PATH",  "Hub Auth",   "No",       "RSA private key PEM path (default: hub_service/.keys/private.pem)."),
              ("HUB_PUBLIC_KEY_PATH",   "Hub Auth",   "No",       "RSA public key PEM path (default: hub_service/.keys/public.pem)."),
              ("HUB_JWKS_URL",          "Hub Auth",   "No",       "Override JWKS URL when hub is behind a load balancer or API gateway."),
              ("AZURE_TENANT_ID",       "Hub Auth",   "If azure", "Azure AD tenant GUID or domain. Required when AUTH_PROVIDER=azure."),
              ("AZURE_CLIENT_ID",       "Hub Auth",   "If azure", "App registration Client ID — expected JWT audience for Entra ID tokens."),
              # MCP server
              ("MCP_AUTH_ENABLED",      "MCP Server", "Yes",      "true in all non-dev environments. false = completely open, no JWT checks."),
              ("MCP_SERVER_ID",         "MCP Server", "Yes",      "Unique server ID matching registry id field. Used as expected JWT audience."),
              ("HUB_SERVER_URL",        "MCP Server", "Yes",      "Hub base URL for JWKS fetch: https://hub.internal:8090"),
              ("MCP_JWT_ISSUER",        "MCP Server", "Yes",      "Expected JWT issuer. Must match HUB_JWT_ISSUER on hub."),
              ("MCP_TOOL_KEY",          "MCP Server", "If ext.",  "API key for MCP-server-to-external-API calls. Never forwarded to agent or hub."),
              # Database
              ("MYSQL_HOST",            "Hub + MCP",  "Yes",      "Database host (default: 127.0.0.1)."),
              ("MYSQL_PORT",            "Hub + MCP",  "Yes",      "Database port (default: 3306)."),
              ("MYSQL_USER",            "Hub + MCP",  "Yes",      "DB username — separate user per service with minimum required privileges."),
              ("MYSQL_PASSWORD",        "Hub + MCP",  "Yes",      "DB password — store in secrets manager, never in source control."),
              ("MYSQL_DATABASE",        "Hub + MCP",  "Yes",      "Target database name (e.g. fab_semantic)."),
              # Chat / infra
              ("CHAT_USERS",            "Chat",       "Yes",      "Comma-separated user:password pairs for initial user seeding."),
              ("OLLAMA_BASE_URL",       "Hub",        "If LLM",   "LLM endpoint for routing agent (default: http://localhost:11434/v1)."),
              ("HUB_JWT_EXPIRY_HOURS",  "Hub",        "No",       "Hub login token expiry in hours (default: 8). Per-server tokens are always 1h."),
          ],
          widths=[3.8, 2.3, 2.0, 7.9])

    doc.add_page_break()

    H2(doc, "6.2  Production Readiness Checklist")
    table(doc,
          ["Area", "Check", "✓"],
          [
              ("Auth",         "AUTH_PROVIDER configured and tested",                            "[ ]"),
              ("Auth",         "JWT_SECRET or HUB_API_KEY set — hub not in dev mode",           "[ ]"),
              ("Auth",         "MCP_AUTH_ENABLED=true on all MCP servers",                       "[ ]"),
              ("Auth",         "MCP_SERVER_ID matches id field in registry",                     "[ ]"),
              ("Auth",         "MCP_JWT_ISSUER matches HUB_JWT_ISSUER",                          "[ ]"),
              ("Auth (Azure)", "AZURE_TENANT_ID and AZURE_CLIENT_ID set if AUTH_PROVIDER=azure", "[ ]"),
              ("Keys",         "RSA key pair exists; private.pem readable only by hub process",  "[ ]"),
              ("Keys",         "HUB_JWT_KID set and matches kid value in JWKS response",        "[ ]"),
              ("Keys",         "/.well-known/jwks.json reachable by all MCP servers via HTTPS", "[ ]"),
              ("Secrets",      ".env and *.pem files excluded from git (.gitignore verified)",   "[ ]"),
              ("Secrets",      "All secrets in vault / secrets manager — not in plaintext files","[ ]"),
              ("Network",      "TLS/HTTPS on all inter-service endpoints",                       "[ ]"),
              ("Network",      "JWKS endpoint served over HTTPS (plain HTTP allows MITM)",       "[ ]"),
              ("Network",      "MCP server ports not publicly reachable — hub-internal only",    "[ ]"),
              ("Database",     "MYSQL_PASSWORD in secrets manager — not in plaintext config",    "[ ]"),
              ("Database",     "hub_events table created and indexed",                           "[ ]"),
              ("Observability","logs/ directory writable by hub process",                        "[ ]"),
              ("Observability","Log aggregator ingesting logs/hub.log (Splunk / ELK / CloudWatch)","[ ]"),
              ("Operations",   "GET /health monitored by infrastructure health check",           "[ ]"),
              ("Operations",   "Rate limiting on /login active and tuned",                       "[ ]"),
          ],
          widths=[2.8, 12.4, 0.8])

    doc.add_page_break()

    H2(doc, "6.3  Security Reference")
    H3(doc, "Security Properties")
    table(doc,
          ["Property", "Implementation", "Threat Mitigated"],
          [
              ("Token forgery prevention",
               "RS256 asymmetric — private key never leaves hub process",
               "Forged tokens rejected at signature verification on every MCP server"),
              ("Cross-server token replay",
               "Per-server aud claim — servers reject mismatched audience",
               "Stolen token from Server A is unusable on Server B"),
              ("Short token exposure window",
               "MCP tokens expire in 1 hour; re-mint via POST /discover",
               "Limits damage duration from interception"),
              ("Credential isolation",
               "Each layer uses its own credential type; none forwarded across boundaries",
               "One compromised component cannot access other layers"),
              ("Password hardening",
               "PBKDF2-SHA256, 200,000 iterations, unique 16-byte salt per user",
               "Offline dictionary attack; rainbow tables"),
              ("Brute-force protection",
               "10 failed logins per username per 15 min → HTTP 423",
               "Online brute-force / credential stuffing"),
              ("PII in logs",
               "audit_log() records argument key names only — never values",
               "Customer data stays out of log streams and SIEM"),
              ("Secrets in version control",
               ".gitignore covers .env · private.pem · public.pem",
               "Credentials exposed in git history"),
              ("Hub dev mode open access",
               "⚠ WARNING at startup; exits when JWT_SECRET or HUB_API_KEY set",
               "Silent open-admin access surviving into production"),
              ("Azure token forgery",
               "PyJWKClient fetches Microsoft JWKS; RS256 sig verified against tenant keys",
               "Forged Entra ID tokens rejected at signature verification"),
          ],
          widths=[4.0, 5.5, 6.5])

    spacer(doc, 4)
    H3(doc, "Anti-Patterns")
    table(doc,
          ["Anti-Pattern", "Risk", "Correct Approach"],
          [
              ("HS256 as sole JWT signing mechanism",
               "Any component with JWT_SECRET can forge tokens",
               "RS256 — servers hold public key only; cannot sign"),
              ("Single JWT valid on all MCP servers",
               "Full cross-service blast radius on token theft",
               "Per-server aud-scoped JWTs from /discover"),
              ("Shared routing agent across requests",
               "Concurrent state interleaving → wrong server selected",
               "New LangGraph agent instance per /discover call"),
              ("Full JWT value in logs",
               "Bearer credential accessible in log streams / SIEM",
               "Log sub and roles from decoded payload only"),
              ("MCP servers publicly reachable",
               "Callers bypass hub auth entirely — skip /discover",
               "Hub-internal network access only"),
              ("AUTH_ENABLED=false in staging",
               "Staging config leaks into production",
               "AUTH_ENABLED=true always; explicitly test auth flow"),
          ],
          widths=[4.8, 4.2, 7.0])

    doc.add_page_break()

    H2(doc, "6.4  Key Files Reference")
    table(doc,
          ["File", "Role", "Notes"],
          [
              ("hub_service/hub_server.py",
               "FastAPI hub — registry, JWT, routing, Admin UI",
               "Entry point; imports auth.py, db.py, observability.py"),
              ("hub_service/auth.py",
               "Hub auth service — token minting, JWKS, multi-provider verification",
               "Exports: verify_token · generate_token · generate_server_token · get_jwks"),
              ("hub_service/db.py",
               "SQLAlchemy engine factory",
               "pool_recycle=1800; .env loaded at import"),
              ("hub_service/observability.py",
               "4-way event fan-out (memory · stdout · file · MySQL)",
               "MySQL permanently disabled on first failure"),
              ("hub_service/.keys/private.pem",
               "RSA private key — signs all JWTs",
               "Never commit to git. Path: HUB_PRIVATE_KEY_PATH"),
              ("hub_service/.keys/public.pem",
               "RSA public key — served at JWKS endpoint",
               "Path: HUB_PUBLIC_KEY_PATH"),
              ("hub_service/mcp-hub.json",
               "Declarative server registry source",
               "Seeded into MySQL via seed_hub_db.py"),
              ("chat_service/chat_server.py",
               "Chat UI server — user auth, streaming, conversation history",
               "PBKDF2 auth; response streaming to browser; asyncio.Task"),
              ("agent.py",
               "LangGraph ReAct orchestrator",
               "Hub discovery; mcp_session(); astream_events(v2)"),
              ("datalayer-as-service/mcp_server/auth.py",
               "MCP-side RBAC — require_role(), audit_log(), BearerClaimsMiddleware",
               "ContextVar-based claim injection; no 2nd JWKS fetch"),
              ("datalayer-as-service/mcp_server/server.py",
               "FastMCP server setup + middleware wiring",
               "JWTVerifier + BearerClaimsMiddleware registration"),
              ("datalayer-as-service/mcp_server/tools.py",
               "MCP tool implementations (@mcp.tool)",
               "Queries semantic views — not raw base tables"),
              ("scripts/seed_hub_db.py",
               "Idempotent DB seeder",
               "UPSERT from mcp-hub.json; safe to re-run"),
              ("datalayer-as-service/.env",
               "MySQL + MCP server secrets",
               "Never commit to git"),
              ("logs/hub.log",
               "JSONL structured event log",
               "Line-buffered; persists across restarts"),
          ],
          widths=[5.2, 5.0, 5.8])


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    print("Generating diagrams...")
    paths = D.generate_all()

    print("Building document...")
    doc = Document()
    _set_page_margins(doc, t=2.0, b=2.0, l=2.0, r=2.0)

    _cover(doc)
    _exec_summary(doc)
    _toc(doc)

    _s1(doc, paths)
    _s2(doc, paths)
    _s3(doc, paths)
    _s4(doc, paths)
    _s5(doc, paths)
    _s6(doc, paths)

    # Footer rule
    spacer(doc, 12)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp,
         f"MCP Hub — Solution Design Document  |  Version 4.0  |  "
         f"{datetime.date.today().strftime('%B %d, %Y')}  |  CONFIDENTIAL",
         italic=True, size=9, color=C["mgrey"])

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
